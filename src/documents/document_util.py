import io
import diff_match_patch as dmp_module
import requests
from docx import Document
import convertapi
import os

convertapi.api_secret = os.environ.get('CONVERT_API_SECRET') or ''


global hashify_url
hashify_url = 'http://api.hashify.net/hash/md5/hex'


# handles creating diffs and giving the hashes
class Document_Util:

    @staticmethod
    def update_document(document, doc_id: str,
                        local_revision_hash: str, content: str) -> str:
        '''
        Given a document_id, local document hash, and new content
        it creates the updated document with the changes from their
        local revision
        '''

        # Make some call to the database to get the content based on the hash
        old_content = document.get_revision_by_hash(
            doc_id, local_revision_hash)

        try:
            if old_content["type"] == "revision":
                old_content = old_content["content"]
            else:
                old_content = ''
        except KeyError:
            old_content = ''

        # Create diff_match_patch object
        dmp = dmp_module.diff_match_patch()

        # set max diff calculate time to 100 milliseconds
        dmp.Diff_Timeout = 0.1

        diff = dmp.diff_main(old_content, content)
        dmp.diff_cleanupSemantic(diff)
        patches = dmp.patch_make(diff)
        current_revision = document.get_most_recent_revision()

        # Apply this patch (changes from their version to new content)
        # to the current revision
        # @TODO: Handle weird behavior with unapplied patches
        updated_text, results = dmp.patch_apply(patches, current_revision)
        return updated_text

    @staticmethod
    def create_hash(content: str) -> str:
        '''Sends the content to hashify to get a hash for the '''
        req = requests.get(f"{hashify_url}?value={content}")
        if 'Digest' in req.json():
            return req.json()['Digest']
        return Exception("There was an error with the Hashify API")

    @staticmethod
    def export_to_pdf(client_id, title, text):
        docx_file = Document_Util.export_to_docx(client_id, title, text)
        params = {
            "FileName": f"{Document_Util._make_safe_filename(title)}",
            "File": convertapi.UploadIO(docx_file, f'{title}.docx'),
            "StoreFile": True,
            "Timeout": 15
        }
        
        result = convertapi.convert('pdf', params, from_format="docx")
        return result.file.url


    @staticmethod
    def export_to_docx(client_id, title, text):
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph(text)
        # safe_client_id = Document_Util._make_safe_filename(client_id)
        # safe_doc_name = Document_Util._make_safe_filename(title)
        # filename = f'{safe_client_id}/{safe_doc_name}.docx'
        file = io.BytesIO()
        document.save(file)
        file.seek(0)
        return file

    def _make_safe_filename(s):
        ''' ensures that filename is safe for filepath, code from stackoverflow'''
        def safe_char(c):
            if c.isalnum():
                return c
            else:
                return "_"
        return "".join(safe_char(c) for c in s).rstrip("_")

